from sqlalchemy.orm import sessionmaker
from models import Tenant, Owner, Property, LeaseContract, Payments, Reporting, engine
from docx import Document
from docx.shared import Pt, Inches
import matplotlib.pyplot as plt
from datetime import datetime
from sqlalchemy import func, and_

Session = sessionmaker(bind=engine)
session = Session()

def generate_word_report(prop_id):
    with Session() as session:
        doc = Document()
        doc.add_heading('Аналіз', level=1)

        report_data = {
            "Звіт по Орендарях": session.query(Tenant).all(),
            "Звіт по Орендодавцях": session.query(Owner).all(),
            "Звіт по Нерухомості": session.query(Property).all(),
            "Звіт по Орендним Договорам": session.query(LeaseContract).all(),
            "Звіт по Платежам": session.query(Payments).all(),
            "Звіт по Аналітичних Даних": session.query(Reporting).all(),
        }

        left_margin = Inches(0.5)
        right_margin = Inches(0.5)

        doc.add_heading('Облік', level=2)

        for title, records in report_data.items():
            doc.add_heading(title, level=3)
            for record in records:
                line = generate_record_line(record)
                add_record_to_doc(doc, line, left_margin, right_margin)

        add_payment_analysis(doc)
        add_occupancy_analysis(doc, property_id=prop_id)
        add_net_profit_analysis(doc, property_id=prop_id)
        add_debt_analysis(doc, property_id=prop_id)

        save_document(doc, prop_id)

def generate_record_line(record):
    if isinstance(record, Tenant):
        return f"Код Орендаря: {record.tenant_id}\nІм'я: {record.name}\nТелефон: {record.phone}\nПошта: {record.email}"
    elif isinstance(record, Owner):
        return f"Код Орендодавця: {record.owner_id}\nІм'я: {record.name}\nТелефон: {record.phone}\nПошта: {record.email}"
    elif isinstance(record, Property):
        return f"Код нерухомості: {record.property_id}\nАдреса: {record.address}\nЦіна: {record.price}\nКод Орендодавця: {record.owner_id}"
    elif isinstance(record, LeaseContract):
        return f"Код Контракту: {record.contract_id}\nПочаток оренди: {record.start_date}\nКінець оренди: {record.end_date}\nКод нерухомості: {record.property_id}\nКод Орендодавця: {record.owner_id}\nКод Орендаря: {record.tenant_id}"
    elif isinstance(record, Payments):
        status = "Оплачено" if record.is_paid else "Не Оплачено"
        return f"Код Платежа: {record.payment_id}\nКод контракту: {record.contract_id}\nДата: {record.date}\nСтатус оплати: {status}"
    elif isinstance(record, Reporting):
        return f"Код звіту: {record.report_id}\nКод нерухомості: {record.property_id}\nКвартал: {record.quarter}\nКількість контрактів: {record.contract_count}\nЗагальний прибуток: {record.total_income}\nЗагальний борг: {record.debt}"

def add_record_to_doc(doc, line, left_margin, right_margin):
    for part in line.split('\n'):
        p = doc.add_paragraph()
        if part:
            key, value = part.split(': ', 1)
            run_key = p.add_run(key + ': ')
            run_key.bold = True
            p.add_run(value)

        p.style.font.size = Pt(13)
        p.paragraph_format.left_indent = left_margin
        p.paragraph_format.right_indent = right_margin

def add_payment_analysis(doc):
    payment_stats = payment_analysis()
    doc.add_heading('Аналіз Платежів', level=2)
    doc.add_paragraph(f"Загальна кількість платежів: {payment_stats['total_payments']}")
    doc.add_paragraph(f"Оплачені платежі: {payment_stats['paid_payments']}")
    doc.add_paragraph(f"Заборговані платежі: {payment_stats['overdue_payments']}")

def add_occupancy_analysis(doc, property_id):
    active_contracts_count = occupancy_rate_analysis(property_id)
    doc.add_heading('Аналіз Заповнюваності', level=2)
    doc.add_paragraph(f"Кількість активних контрактів для об’єкта {property_id}: {active_contracts_count}")

def add_net_profit_analysis(doc, property_id):
    net_profit = calculate_net_profit(property_id)
    doc.add_heading('Чистий Прибуток', level=2)
    doc.add_paragraph(f"Чистий прибуток для об'єкта {property_id}: {net_profit}")

def add_debt_analysis(doc, property_id):
    total_debt = debt_analysis(property_id)
    doc.add_heading('Аналіз Заборгованості', level=2)
    doc.add_paragraph(f"Загальна заборгованість для об'єкта {property_id}: {total_debt}")

def save_document(doc, prop_id):
    doc_file = f"full_report_property_{prop_id}.docx"
    try:
        doc.save(doc_file)
        print(f"Звіт згенеровано: {doc_file}")
    except PermissionError:
        print(f"Модуль не може бути виконаний поки відкритий '{doc_file}'")

def calculate_net_profit(property_id):
    total_income = session.query(func.sum(Property.price)).join(LeaseContract).join(Payments).filter(
        and_(
            LeaseContract.property_id == property_id,
            Payments.is_paid == 1
        )
    ).scalar() or 0
    return total_income



def payment_analysis():
    total_payments = session.query(Payments).count()
    paid_payments = session.query(Payments).filter(Payments.is_paid == 1).count()
    overdue_payments = session.query(Payments).filter(Payments.is_paid == 0).count()
    return {
        "total_payments": total_payments,
        "paid_payments": paid_payments,
        "overdue_payments": overdue_payments,
    }

def occupancy_rate_analysis(property_id):
    active_contracts = session.query(LeaseContract).filter(
        LeaseContract.property_id == property_id,
        LeaseContract.end_date >= datetime.now()
    ).count()
    return active_contracts

def debt_analysis(property_id):
    total_debt = session.query(func.sum(Property.price)).join(LeaseContract).join(Payments).filter(
        and_(
            LeaseContract.property_id == property_id,
            Payments.is_paid == 0
        )
    ).scalar() or 0
    return total_debt

def plot_income_report():
    income_data = (
        session.query(
            Property.address,
            Property.price,
            func.count(Payments.payment_id).label('paid_payments_count')
        )
        .outerjoin(LeaseContract, LeaseContract.property_id == Property.property_id)
        .outerjoin(Payments, Payments.contract_id == LeaseContract.contract_id)
        .filter(Payments.is_paid == 1)
        .group_by(Property.address, Property.price)
        .all()
    )

    addresses = [data[0] for data in income_data]
    prices = [data[1] for data in income_data]
    paid_counts = [data[2] for data in income_data]
    total_incomes = [price * count for price, count in zip(prices, paid_counts)]

    plt.figure(figsize=(10, 6))
    plt.bar(addresses, total_incomes, color='skyblue')
    plt.xlabel('Об\'єкти нерухомості')
    plt.ylabel('Загальний дохід')
    plt.title('Загальний дохід по об\'єктах нерухомості')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()

